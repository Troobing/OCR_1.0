"""
Word 文档生成 — 将 AI 提取内容渲染为 .docx
负责：内容解析、LaTeX → OMML 转换、文档拼装
Config: 字体/字号/页边距、公式转换 XSLT、合并/分开下载逻辑
（所有端共用 — HTTP 和 Bridge 都调它）
Skill：python-docx、XSLT、LaTeX→MathML→OMML、Markdown 表格解析

处理流水线：
  解析内容(分离文字/公式/表格) → LaTeX → MathML → XSLT → OMML → python-docx 拼装 .docx
"""

import re
import io
import zipfile
from typing import Optional

from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# OMML 命名空间（手工构造矩阵/方程组时使用）
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# ─── MathML → OMML XSLT 转换表 ───

_MATHML_TO_OMML_XSLT = r'''<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
  xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
  xmlns:math="http://www.w3.org/1998/Math/MathML"
  xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
  exclude-result-prefixes="math">

  <xsl:template match="text()"><xsl:value-of select="."/></xsl:template>
  <xsl:template match="math:math"><m:oMath><xsl:apply-templates/></m:oMath></xsl:template>

  <xsl:template match="math:mi">
    <m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t><xsl:apply-templates/></m:t></m:r>
  </xsl:template>
  <xsl:template match="math:mn">
    <m:r><m:t><xsl:apply-templates/></m:t></m:r>
  </xsl:template>
  <xsl:template match="math:mo">
    <m:r><m:t><xsl:apply-templates/></m:t></m:r>
  </xsl:template>
  <xsl:template match="math:mtext">
    <m:r><m:t><xsl:apply-templates/></m:t></m:r>
  </xsl:template>
  <xsl:template match="math:mrow"><xsl:apply-templates/></xsl:template>

  <xsl:template match="math:msup">
    <m:sSup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sup><xsl:apply-templates select="*[2]"/></m:sup>
    </m:sSup>
  </xsl:template>
  <xsl:template match="math:msub">
    <m:sSub>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
    </m:sSub>
  </xsl:template>
  <xsl:template match="math:msubsup">
    <m:sSubSup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
      <m:sup><xsl:apply-templates select="*[3]"/></m:sup>
    </m:sSubSup>
  </xsl:template>

  <xsl:template match="math:mfrac">
    <m:f>
      <m:fPr/>
      <m:num><xsl:apply-templates select="*[1]"/></m:num>
      <m:den><xsl:apply-templates select="*[2]"/></m:den>
    </m:f>
  </xsl:template>

  <xsl:template match="math:msqrt">
    <m:rad><m:radPr/><m:e><xsl:apply-templates/></m:e></m:rad>
  </xsl:template>
  <xsl:template match="math:mroot">
    <m:rad>
      <m:radPr><m:degHide m:val="0"/></m:radPr>
      <m:deg><xsl:apply-templates select="*[2]"/></m:deg>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
    </m:rad>
  </xsl:template>

  <xsl:template match="math:mfenced">
    <m:d>
      <m:dPr>
        <m:begChr m:val="{@open}"/>
        <m:endChr m:val="{@close}"/>
      </m:dPr>
      <m:e><xsl:apply-templates/></m:e>
    </m:d>
  </xsl:template>

  <xsl:template match="math:mover">
    <m:acc>
      <m:accPr><m:chr m:val="{*[2]/text()}"/></m:accPr>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
    </m:acc>
  </xsl:template>

  <!-- mstyle 透传：救 \dfrac \tfrac \cfrac \dbinom \tbinom（外层 mstyle 无意义，直接展开内部） -->
  <xsl:template match="math:mstyle"><xsl:apply-templates/></xsl:template>

  <!-- munder → m:limLow：救 \underline \underset \underbrace（下方修饰） -->
  <xsl:template match="math:munder">
    <m:limLow>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:lim><xsl:apply-templates select="*[2]"/></m:lim>
    </m:limLow>
  </xsl:template>

  <!-- mspace 透传忽略：行分隔由 Python 预处理负责（救 aligned 残留的 mspace linebreak） -->
  <xsl:template match="math:mspace"/>

  <!-- mpadded 透传：救 \not= \xrightarrow 等组合命令的外层容器 -->
  <xsl:template match="math:mpadded"><xsl:apply-templates/></xsl:template>

  <xsl:template match="math:*">
    <m:r><m:t><xsl:value-of select="."/></m:t></m:r>
  </xsl:template>

</xsl:stylesheet>
'''

# ─── XSLT 懒加载编译 ───

_xslt_transform: Optional[etree.XSLT] = None

def _get_xslt() -> etree.XSLT:
    global _xslt_transform
    if _xslt_transform is None:
        _xslt_transform = etree.XSLT(
            etree.fromstring(_MATHML_TO_OMML_XSLT.encode("utf-8"))
        )
    return _xslt_transform

# ─── 主入口 ───

def generate_word(
    contents: list[str],
    image_ids: list[str],
    merge: bool = True,
) -> tuple[bytes, str, str]:
    """生成 Word 文档。merge=True → 单个 .docx，merge=False → .zip 内含多个 .docx"""
    if merge:
        doc = Document()
        _setup_document_style(doc)
        for i, content in enumerate(contents):
            if i > 0:
                doc.add_page_break()
            _build_document_from_content(doc, content)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), "提取结果.docx", \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for content, image_id in zip(contents, image_ids):
                doc = Document()
                _setup_document_style(doc)
                _build_document_from_content(doc, content)
                file_buffer = io.BytesIO()
                doc.save(file_buffer)
                file_buffer.seek(0)
                zf.writestr(f"{image_id}.docx", file_buffer.getvalue())
        zip_buffer.seek(0)
        return zip_buffer.getvalue(), "提取结果.zip", "application/zip"

# ─── 文档样式 ───

def _setup_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # 设置中文字体槽
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    rfonts.set(qn('w:eastAsia'), '宋体')

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ─── 内容构建 ───

def _build_document_from_content(doc: Document, content: str) -> None:
    """解析 content，将文字+公式+表格混合内容渲染到 doc 中"""
    segments = _parse_content(content)
    groups = _group_segments(segments)
    for group in groups:
        if len(group) == 1 and group[0]["type"] == "block_formula":
            omml = _latex_to_omml(group[0]["content"])
            if omml is not None:
                _add_omml_paragraph(doc, omml, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _add_fallback_formula(doc, group[0]["content"], block=True)
        elif len(group) == 1 and group[0]["type"] == "table":
            _add_markdown_table(doc, group[0]["content"])
        else:
            _add_mixed_paragraph(doc, group)

def _group_segments(segments: list[dict]) -> list[list[dict]]:
    """将 text + inline_formula 归为一组，block_formula 和 table 单独成组"""
    groups: list[list[dict]] = []
    current: list[dict] = []
    for seg in segments:
        if seg["type"] in ("block_formula", "table"):
            if current:
                groups.append(current)
                current = []
            groups.append([seg])
        else:
            current.append(seg)
    if current:
        groups.append(current)
    return groups

def _add_mixed_paragraph(doc: Document, group: list[dict]) -> None:
    """一行内的文字+行内公式放在同一个 Word 段落中，避免公式强制换行"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    for seg in group:
        if seg["type"] == "text":
            lines = _sanitize_xml_text(seg["content"]).split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    br_run = OxmlElement("w:r")
                    br_run.append(OxmlElement("w:br"))
                    para._element.append(br_run)
                if line.strip():
                    para.add_run(line)
        elif seg["type"] == "inline_formula":
            omml = _latex_to_omml(seg["content"])
            if omml is not None:
                try:
                    omml_elem = etree.fromstring(omml.encode("utf-8"))
                    run_elem = OxmlElement("w:r")
                    run_elem.append(omml_elem)
                    para._element.append(run_elem)
                except Exception:
                    para.add_run(f"${_sanitize_xml_text(seg['content'])}$")
            else:
                fallback_run = para.add_run(f"${_sanitize_xml_text(seg['content'])}$")
                fallback_run.font.name = "Consolas"
                fallback_run.font.size = Pt(10.5)

# ─── Markdown 表格渲染 ───

def _add_markdown_table(doc: Document, table_text: str) -> None:
    """将 Markdown 表格文本渲染为 Word 表格"""
    lines = [ln.strip() for ln in _sanitize_xml_text(table_text).strip().split("\n") if ln.strip()]
    if len(lines) < 2:
        # 不够构成表格，降级为纯文本
        para = doc.add_paragraph()
        para.add_run(_sanitize_xml_text(table_text))
        return

    # 第 2 行是分隔行（|---|---|），用于判定
    def _is_separator(line: str) -> bool:
        return bool(re.match(r"^\|?[\s:|-]+\|?$", line)) and "---" in line

    header = _split_table_row(lines[0])
    body_rows: list[list[str]] = []
    start = 1
    if len(lines) >= 2 and _is_separator(lines[1]):
        start = 2
    for ln in lines[start:]:
        body_rows.append(_split_table_row(ln))

    cols = max(len(header), max((len(r) for r in body_rows), default=0))
    if cols == 0:
        para = doc.add_paragraph()
        para.add_run(_sanitize_xml_text(table_text))
        return

    # 补齐列数
    header = header + [""] * (cols - len(header))
    body_rows = [r + [""] * (cols - len(r)) for r in body_rows]

    table = doc.add_table(rows=1 + len(body_rows), cols=cols)
    table.style = "Table Grid"
    # 表头
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = _sanitize_xml_text(cell_text)
        # 表头加粗
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # 数据行
    for i, row in enumerate(body_rows, start=1):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = _sanitize_xml_text(cell_text)

def _split_table_row(line: str) -> list[str]:
    """拆分 Markdown 表格行：去掉首尾 |，按 | 切分"""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]

# ─── 正则解析：拆分文字 & 公式 & 表格 ───

# 公式匹配：行内禁内部 $ 和换行，块公式用否定向前查找避免提前闭合
_FORMULA_PATTERN = re.compile(
    r"(\$\$((?:(?!\$\$).)+?)\$\$|\$([^$\n]+?)\$)",
    re.DOTALL,
)
# Markdown 表格：连续两行以上以 | 开头的行（含分隔行）
_TABLE_LINE = re.compile(r"^\s*\|.*\|\s*$")

# XML 不兼容的控制字符（保留 \t \n \r）—— LLM 偶发输出退格/换页等会令 python-docx 崩溃
_XML_CTRL_STRIP = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _sanitize_xml_text(text: str) -> str:
    """剥离 XML 不兼容的控制字符，避免 python-docx 在 add_run 时抛 ValueError。"""
    if not text:
        return text
    return _XML_CTRL_STRIP.sub("", text)


def _parse_content(content: str) -> list[dict]:
    """用正则将混合文字和 $...$/$$...$$ 与 Markdown 表格的内容拆成结构化片段。

    处理顺序：
      1. 先把 \\$ 转义占位，避免金额符号被误当公式边界
      2. 扫描表格块（连续 |...| 行）
      3. 扫描公式
      4. 还原转义占位
    """
    # 1. 转义 \$ 占位
    escapes: list[str] = []

    def _stash_esc(m: re.Match) -> str:
        escapes.append(m.group(0))
        return f"\x00ESC{len(escapes) - 1}\x00"

    content = re.sub(r"\\\$", _stash_esc, content)

    # 2. + 3. 先按行扫描表格块，把表格行整体替换为占位，再让公式正则处理剩余文本
    # 这样表格内的 $ 不会被当公式
    tables: list[str] = []
    lines = content.split("\n")
    out_lines: list[str] = []
    i = 0
    while i < len(lines):
        if _TABLE_LINE.match(lines[i]):
            # 收集连续表格行
            block_start = i
            while i < len(lines) and _TABLE_LINE.match(lines[i]):
                i += 1
            table_text = "\n".join(lines[block_start:i])
            tables.append(table_text)
            out_lines.append(f"\x00TBL{len(tables) - 1}\x00")
        else:
            out_lines.append(lines[i])
            i += 1
    content = "\n".join(out_lines)

    # 4. 公式匹配
    segments: list[dict] = []
    pos = 0
    for match in _FORMULA_PATTERN.finditer(content):
        if match.start() > pos:
            text_before = content[pos:match.start()]
            if text_before.strip():
                segments.append({"type": "text", "content": _restore_escapes(text_before, escapes)})
        if match.group(2) is not None:
            segments.append({"type": "block_formula", "content": match.group(2).strip()})
        else:
            segments.append({"type": "inline_formula", "content": match.group(3).strip()})
        pos = match.end()
    if pos < len(content):
        remaining = content[pos:]
        if remaining.strip():
            segments.append({"type": "text", "content": _restore_escapes(remaining, escapes)})

    # 5. 把表格占位替换回 table segment（按出现顺序插入）
    result: list[dict] = []
    tbl_idx = 0
    for seg in segments:
        if seg["type"] == "text":
            text = seg["content"]
            # 文本中可能含 \x00TBLn\x00 占位
            parts = re.split(r"\x00TBL(\d+)\x00", text)
            for k, part in enumerate(parts):
                if k % 2 == 0:
                    if part.strip():
                        result.append({"type": "text", "content": part})
                else:
                    idx = int(part)
                    if idx < len(tables):
                        result.append({"type": "table", "content": tables[idx]})
        else:
            result.append(seg)
    return result


def _restore_escapes(text: str, escapes: list[str]) -> str:
    r"""还原 \$ 转义占位"""
    if not escapes:
        return text

    def _restore(m: re.Match) -> str:
        idx = int(m.group(1))
        if 0 <= idx < len(escapes):
            return escapes[idx]
        return m.group(0)

    return re.sub(r"\x00ESC(\d+)\x00", _restore, text)

# ─── LaTeX → OMML 转换 ───

# 环境匹配：识别 \begin{env} body \end{env}（DOTALL 让 . 跨行）
_ENV_FULLMATCH = re.compile(r"^\\begin\{(\w+)\}(.*?)\\end\{\1\}\s*$", re.DOTALL)

# 环境到定界符的映射（OMML m:d 的 begChr/endChr）
_ENV_DELIMS = {
    "pmatrix":     ("(", ")"),
    "bmatrix":     ("[", "]"),
    "Bmatrix":     ("{", "}"),
    "vmatrix":     ("|", "|"),
    "Vmatrix":     ("\u2016", "\u2016"),   # ‖
    "matrix":      ("", ""),
    "smallmatrix": ("", ""),
    "array":       ("", ""),
    "cases":       ("{", ""),
}

_EQARR_ENVS = {"aligned", "align", "gathered", "gather"}


def _xml_escape(text: str) -> str:
    """XML 文本节点转义"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _split_env_body(body: str) -> list[list[str]]:
    """把环境体按 \\ 拆行、按 & 拆列。

    使用负向零宽断言避开转义的 \\\\ 与 \\&。
    返回二维列表，每个元素是单元格的 LaTeX 片段（未 trim，由调用方决定）。
    """
    # 行分隔符是 \\（两个反斜杠）；正则里写 \\\\\\\\ 表示匹配字面 \\
    rows = re.split(r"(?<!\\)\\\\", body)
    return [re.split(r"(?<!\\)&", r) for r in rows]


def _latex_to_omml_inner(latex: str) -> str:
    """单元格内部公式转换：不走环境预处理，避免无限递归。

    成功返回 OMML 片段（不含 m:oMath 根），失败返回纯文本 m:r。
    """
    omml = _latex_to_omml_via_mathml(latex)
    if omml is None:
        # 降级：纯文本 run
        return f'<m:r><m:t>{_xml_escape(latex)}</m:t></m:r>'
    # 剥掉 <?xml ...?> 头和 m:oMath 根，只取内部
    # _latex_to_omml_via_mathml 返回形如 '<?xml version="1.0"?>\\n<m:oMath ...>INNER</m:oMath>'
    m = re.search(r"<m:oMath[^>]*>(.*)</m:oMath>", omml, re.DOTALL)
    if m:
        return m.group(1)
    return omml


def _latex_to_omml_via_mathml(latex: str) -> Optional[str]:
    """走 latex2mathml → XSLT 的原链路。失败返回 None。"""
    try:
        from latex2mathml import converter
        mathml_str = converter.convert(latex)
        mathml_elem = etree.fromstring(mathml_str.encode("utf-8"))
        transform = _get_xslt()
        return str(transform(mathml_elem))
    except Exception:
        return None


def _build_omml_matrix(rows: list[list[str]], open_d: str, close_d: str) -> str:
    """构造 OMML 矩阵 XML 片段（不含 m:oMath 根）。

    - rows: 二维 LaTeX 字符串列表
    - open_d/close_d: 定界符（空串表示无）
    列数取所有行的最大值，不足补空。
    """
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    if cols == 0:
        cols = 1

    # 列规格：每列居中对齐
    mc_list = "\n".join(
        f'        <m:mc><m:mcPr><m:count m:val="1"/><m:mcJc m:val="center"/></m:mcPr></m:mc>'
        for _ in range(cols)
    )
    # 行
    mr_list = []
    for r in rows:
        # 补齐列
        cells = list(r) + [""] * (cols - len(r))
        e_list = "\n".join(
            f"          <m:e>{_latex_to_omml_inner(c.strip())}</m:e>"
            for c in cells
        )
        mr_list.append(
            f"        <m:mr>\n{e_list}\n        </m:mr>"
        )

    matrix_xml = (
        f"      <m:m>\n"
        f"        <m:mPr>\n"
        f"          <m:mcs>\n{mc_list}\n          </m:mcs>\n"
        f"        </m:mPr>\n"
        f"{chr(10).join(mr_list)}\n"
        f"      </m:m>"
    )

    if not open_d and not close_d:
        return matrix_xml

    # 包定界符
    return (
        f"      <m:d>\n"
        f"        <m:dPr><m:begChr m:val=\"{_xml_escape(open_d)}\"/>"
        f"<m:endChr m:val=\"{_xml_escape(close_d)}\"/></m:dPr>\n"
        f"        <m:e>\n{matrix_xml}\n        </m:e>\n"
        f"      </m:d>"
    )


def _build_omml_eqarr(rows: list[list[str]]) -> str:
    """构造 OMML 等式数组（m:eqArr，对应 aligned/gathered）。

    aligned 的 & 仅作对齐参考——把每行的多个单元格 join 成一行公式，
    用空格替代 &（OMML 的 eqArr 不支持列对齐，只支持逐行显示）。
    """
    if not rows:
        return ""
    e_list = []
    for r in rows:
        # 行内多列用空格连接（丢对齐但保留可读）
        row_latex = " ".join(c.strip() for c in r if c.strip())
        e_list.append(f"        <m:e>{_latex_to_omml_inner(row_latex)}</m:e>")
    return (
        f"      <m:eqArr>\n{chr(10).join(e_list)}\n      </m:eqArr>"
    )


def _build_env_omml(env_name: str, body: str) -> Optional[str]:
    """环境入口：根据 env_name 分发到 matrix / eqarr 构造器。

    返回 OMML 片段（不含 m:oMath 根），失败返回 None（调用方 fall through）。
    """
    try:
        if env_name in _EQARR_ENVS:
            rows = _split_env_body(body)
            return _build_omml_eqarr(rows)
        if env_name in _ENV_DELIMS:
            # array 有列规格 {|c|c|}，先剥掉
            if env_name == "array":
                body = re.sub(r"^\s*\{[^}]*\}", "", body, count=1)
            rows = _split_env_body(body)
            open_d, close_d = _ENV_DELIMS[env_name]
            return _build_omml_matrix(rows, open_d, close_d)
        return None
    except Exception:
        return None


def _latex_to_omml(latex: str) -> Optional[str]:
    """LaTeX → OMML。失败返回 None，调用方降级为纯文本。

    处理顺序：
      1. 矩阵/方程组/aligned 等环境 → Python 预处理直接构造 OMML
         （绕过 latex2mathml 对环境的多行压制 bug 和 aligned 的 XML 崩溃）
      2. 其他公式 → latex2mathml → XSLT → OMML
      3. 任一失败返回 None
    """
    if not latex:
        return None
    latex = latex.strip()

    # 1. 环境预处理
    m = _ENV_FULLMATCH.match(latex)
    if m:
        env_omml = _build_env_omml(m.group(1), m.group(2))
        if env_omml is not None:
            return f'<m:oMath xmlns:m="{_M_NS}">{env_omml}</m:oMath>'
        # 失败则 fall through 到 latex2mathml

    # 2. 普通公式走原链路
    return _latex_to_omml_via_mathml(latex)

# ─── 辅助：向 Document 添加内容块 ───

def _add_omml_paragraph(doc: Document, omml_xml: str, alignment: Optional[int] = None) -> None:
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    try:
        omml_elem = etree.fromstring(omml_xml.encode("utf-8"))
        run_elem = OxmlElement("w:r")
        run_elem.append(omml_elem)
        para._element.append(run_elem)
    except etree.XMLSyntaxError as e:
        para.add_run(f"[公式解析错误: {e}]")

def _add_fallback_formula(doc: Document, latex: str, block: bool = False) -> None:
    para = doc.add_paragraph()
    if block:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(_sanitize_xml_text(latex))
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)
